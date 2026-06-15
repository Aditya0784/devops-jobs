"""Backend API tests for Resume Tailor."""
import io
import os
import pytest
import requests
from pypdf import PdfWriter
from docx import Document

BASE_URL = os.environ.get("REACT_APP_BACKEND_URL", "http://localhost:8000").rstrip("/")
API = f"{BASE_URL}/api"

SCRAPE_SLUGS = ["stripe", "airbnb", "databricks", "gitlab", "cloudflare", "netflix"]


@pytest.fixture(scope="session")
def session():
    s = requests.Session()
    return s


@pytest.fixture(scope="session")
def scraped(session):
    r = session.post(f"{API}/scrape", json={"slugs": SCRAPE_SLUGS}, timeout=120)
    assert r.status_code == 200, r.text
    return r.json()


# --- Health / metadata ---
def test_root(session):
    r = session.get(f"{API}/", timeout=30)
    assert r.status_code == 200
    data = r.json()
    assert data.get("status") == "ok"
    assert data.get("app") == "resume-tailor"


def test_companies(session):
    r = session.get(f"{API}/companies", timeout=30)
    assert r.status_code == 200
    data = r.json()
    assert data["count"] > 0
    assert isinstance(data["companies"], list)
    assert "slug" in data["companies"][0]


# --- Scrape ---
def test_scrape(scraped):
    assert scraped["count"] > 0, f"Scrape returned 0 jobs: {scraped}"


# --- Jobs listing ---
def test_jobs_list(scraped, session):
    r = session.get(f"{API}/jobs", timeout=30)
    assert r.status_code == 200
    data = r.json()
    assert data["count"] > 0
    valid_roles = {"devops", "sre", "cloud_architect"}
    for j in data["jobs"]:
        assert j["role_type"] in valid_roles, j


def test_jobs_filter_role_sre(scraped, session):
    r = session.get(f"{API}/jobs", params={"role": "sre"}, timeout=30)
    assert r.status_code == 200
    data = r.json()
    for j in data["jobs"]:
        assert j["role_type"] == "sre"


def test_jobs_search_kubernetes(scraped, session):
    r = session.get(f"{API}/jobs", params={"q": "kubernetes"}, timeout=30)
    assert r.status_code == 200
    data = r.json()
    # search is case-insensitive; results may be 0 or more, but request should succeed
    for j in data["jobs"]:
        haystack = (j.get("title", "") + " " + j.get("description", "") + " " + j.get("location", "")).lower()
        assert "kubernetes" in haystack


def test_get_job_by_id(scraped, session):
    r = session.get(f"{API}/jobs", timeout=30)
    jobs = r.json()["jobs"]
    assert jobs
    jid = jobs[0]["id"]
    r2 = session.get(f"{API}/jobs/{jid}", timeout=30)
    assert r2.status_code == 200
    assert r2.json()["id"] == jid


def test_get_job_404(session):
    r = session.get(f"{API}/jobs/nonexistent-id-xyz", timeout=30)
    assert r.status_code == 404


# --- Resume upload ---
def _make_pdf_bytes() -> bytes:
    # Generate a real PDF using pypdf with one blank page (resume_parser must extract some text or empty)
    # pypdf needs text via reportlab-like approach; instead, use a small known-good PDF text via fpdf if available.
    # Fall back: try fpdf2
    try:
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.cell(0, 10, "John Doe - DevOps Engineer with Kubernetes and AWS experience.")
        pdf.ln()
        pdf.cell(0, 10, "Skills: Docker, Terraform, CI/CD, Python, Linux.")
        out = pdf.output(dest="S")
        if isinstance(out, str):
            return out.encode("latin-1")
        return bytes(out)
    except Exception:
        # Last resort: empty PDF (will fail extraction)
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        buf = io.BytesIO()
        writer.write(buf)
        return buf.getvalue()


def _make_docx_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("Jane Doe - SRE with strong Kubernetes, Prometheus, and AWS background.")
    doc.add_paragraph("Skills: Python, Go, Terraform, GCP.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture(scope="session")
def resume_id(session):
    data = _make_pdf_bytes()
    files = {"file": ("resume.pdf", data, "application/pdf")}
    r = session.post(f"{API}/resume/upload", files=files, timeout=60)
    assert r.status_code == 200, r.text
    j = r.json()
    assert "id" in j and j["chars"] > 0 and "preview" in j
    return j["id"]


def test_upload_pdf(resume_id):
    assert resume_id


def test_upload_docx(session):
    data = _make_docx_bytes()
    files = {"file": ("resume.docx", data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    r = session.post(f"{API}/resume/upload", files=files, timeout=60)
    assert r.status_code == 200, r.text
    j = r.json()
    assert j["chars"] > 0


def test_upload_txt_rejected(session):
    files = {"file": ("resume.txt", b"hello world", "text/plain")}
    r = session.post(f"{API}/resume/upload", files=files, timeout=30)
    assert r.status_code == 400


# --- Analyze error paths ---
def test_analyze_404_resume(scraped, session):
    r = session.get(f"{API}/jobs", timeout=30)
    jid = r.json()["jobs"][0]["id"]
    body = {"resume_id": "bad-resume-id", "job_id": jid, "provider": "openai", "api_key": "sk-thisIsAFakeKey12345"}
    r2 = session.post(f"{API}/analyze", json=body, timeout=30)
    assert r2.status_code == 404


def test_analyze_404_job(session, resume_id):
    body = {"resume_id": resume_id, "job_id": "bad-job-id", "provider": "openai", "api_key": "sk-thisIsAFakeKey12345"}
    r = session.post(f"{API}/analyze", json=body, timeout=30)
    assert r.status_code == 404


def test_analyze_400_no_key(session, resume_id, scraped):
    r = session.get(f"{API}/jobs", timeout=30)
    jid = r.json()["jobs"][0]["id"]
    body = {"resume_id": resume_id, "job_id": jid, "provider": "openai", "api_key": "short"}
    r2 = session.post(f"{API}/analyze", json=body, timeout=30)
    assert r2.status_code == 400


# --- Download error path ---
def test_download_404(session):
    r = session.post(f"{API}/download", json={"analysis_id": "nope", "format": "pdf"}, timeout=30)
    assert r.status_code == 404
