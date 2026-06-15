"""Parse resumes from PDF or DOCX files."""

import io
from pypdf import PdfReader
from docx import Document


def parse_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    chunks = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text:
            chunks.append(text)
    return "\n".join(chunks).strip()


def parse_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    chunks = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    chunks.append(cell.text)
    return "\n".join(chunks).strip()


def parse_resume(file_bytes: bytes, filename: str) -> str:
    fn = (filename or "").lower()
    if fn.endswith(".pdf"):
        return parse_pdf(file_bytes)
    if fn.endswith(".docx"):
        return parse_docx(file_bytes)
    raise ValueError("Unsupported file type. Upload .pdf or .docx")
