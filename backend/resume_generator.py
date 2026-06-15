"""Generate tailored resume as downloadable PDF or DOCX from plain text."""

import io
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document
from docx.shared import Pt


SECTION_HEADERS = {
    "SUMMARY", "PROFILE", "OBJECTIVE",
    "SKILLS", "TECHNICAL SKILLS", "CORE SKILLS",
    "EXPERIENCE", "WORK EXPERIENCE", "PROFESSIONAL EXPERIENCE", "EMPLOYMENT",
    "PROJECTS", "KEY PROJECTS",
    "EDUCATION",
    "CERTIFICATIONS", "CERTIFICATES",
    "ACHIEVEMENTS", "AWARDS",
}


def _is_header(line: str) -> bool:
    s = line.strip().rstrip(":").upper()
    return s in SECTION_HEADERS


def to_pdf(text: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
        topMargin=0.6 * inch, bottomMargin=0.6 * inch,
    )
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=14, spaceAfter=6, alignment=TA_LEFT)
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=10, leading=13)
    story = []
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            story.append(Spacer(1, 6))
            continue
        if _is_header(line):
            story.append(Spacer(1, 6))
            story.append(Paragraph(line.strip().rstrip(":").upper(), h1))
        else:
            safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe, body))
    doc.build(story)
    return buf.getvalue()


def to_docx(text: str) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        if _is_header(line):
            p = doc.add_paragraph()
            run = p.add_run(line.strip().rstrip(":").upper())
            run.bold = True
            run.font.size = Pt(13)
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
